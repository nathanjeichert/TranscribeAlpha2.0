import io
import logging
import os
import re
from typing import List, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt

try:
    from .models import TranscriptTurn, WordTimestamp
except ImportError:
    try:
        from models import TranscriptTurn, WordTimestamp
    except ImportError:
        import models
        TranscriptTurn = models.TranscriptTurn
        WordTimestamp = models.WordTimestamp

logger = logging.getLogger(__name__)


# Shared layout constants used for XML generation and editor exports
SPEAKER_PREFIX_SPACES = 10  # Leading spaces before speaker name in XML (visual simulation)
CONTINUATION_SPACES = 0     # Leading spaces for continuation lines in XML (visual simulation)
SPEAKER_COLON = ":   "      # Colon and spaces after speaker name (total 4 chars)
MAX_TOTAL_LINE_WIDTH = 64   # Maximum total characters per XML line for speaker lines
MAX_CONTINUATION_WIDTH = 64 # Maximum total characters per XML line for continuation lines
MIN_LINE_DURATION_SECONDS = 1.25

# OnCue XML constants
ONCUE_FIRST_PGLN = 101      # First page-line number (page 1, line 1 = 101)
DEFAULT_VIDEO_ID = "1"      # Default video ID for single-video transcripts


def timestamp_to_seconds(timestamp: Optional[str]) -> float:
    """Convert timestamp like '[MM:SS]' or 'MM:SS' to seconds."""
    if not timestamp:
        return 0.0
    ts = timestamp.strip('[]').strip()
    parts = ts.split(':')
    try:
        if len(parts) == 3:
            h, m, s = map(float, parts)
            return h * 3600 + m * 60 + s
        if len(parts) == 2:
            m, s = map(float, parts)
            return m * 60 + s
        return float(ts)
    except ValueError:
        return 0.0


def seconds_to_timestamp(seconds: float) -> str:
    """Convert seconds float to OnCue-style [MM:SS] or [HH:MM:SS] timestamp."""
    if seconds < 0:
        seconds = 0.0
    total_seconds = int(round(seconds))
    hours, remainder = divmod(total_seconds, 3600)
    minutes, secs = divmod(remainder, 60)
    if hours > 0:
        return f"[{hours:02d}:{minutes:02d}:{secs:02d}]"
    return f"[{minutes:02d}:{secs:02d}]"


def wrap_text_for_transcript(text: str, max_width: int) -> List[str]:
    """
    Wrap text to fit within max_width characters, preserving word boundaries.

    Args:
        text: The text to wrap
        max_width: Maximum characters per line of text content

    Returns:
        List of wrapped lines
    """
    if not text:
        return [""]

    if max_width <= 0:
        return [text]

    words = text.split()
    lines = []
    current_line = []
    current_length = 0

    for word in words:
        word_length = len(word)
        # +1 for space before word (except first word)
        space_needed = word_length + (1 if current_line else 0)

        if current_length + space_needed <= max_width:
            current_line.append(word)
            current_length += space_needed
        else:
            if current_line:
                lines.append(" ".join(current_line))
            current_line = [word]
            current_length = word_length

    if current_line:
        lines.append(" ".join(current_line))

    return lines if lines else [""]


def create_docx(title_data: dict, transcript_turns: List[TranscriptTurn]) -> bytes:
    """
    Generate DOCX transcript from transcript turns and title data.
    """
    # Create a new document
    doc = Document()

    # Set margins and page layout
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Add title page
    title_lines = [
        "Generated Transcript",
        "",
        f"Case Name: {title_data.get('CASE_NAME', '')}",
        f"Case Number: {title_data.get('CASE_NUMBER', '')}",
        f"Date: {title_data.get('DATE', '')}",
        f"Time: {title_data.get('TIME', '')}",
        f"Location: {title_data.get('LOCATION', '')}",
        f"Original File: {title_data.get('FILE_NAME', '')}",
        f"Duration: {title_data.get('FILE_DURATION', '')}",
        f"Firm/Organization: {title_data.get('FIRM_OR_ORGANIZATION_NAME', '')}",
    ]

    for line in title_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 2.0

    doc.add_page_break()

    # Add transcript turns
    if transcript_turns:
        # Remove extra empty paragraph after page break
        p_element = doc.paragraphs[-1]._element
        p_element.getparent().remove(p_element)
        for turn in transcript_turns:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.first_line_indent = Inches(1.0)  # Standard legal transcript indent
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.widow_control = False

            # Only show speaker label if not a continuation turn
            if not turn.is_continuation:
                speaker_run = p.add_run(f"{turn.speaker.upper()}:   ")
                speaker_run.font.name = "Courier New"
            text_run = p.add_run(turn.text)
            text_run.font.name = "Courier New"
    else:
        for turn in transcript_turns:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.first_line_indent = Inches(1.0)
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.widow_control = False

            # Only show speaker label if not a continuation turn
            if not turn.is_continuation:
                speaker_run = p.add_run(f"{turn.speaker.upper()}:   ")
                speaker_run.font.name = "Courier New"
            text_run = p.add_run(turn.text)
            text_run.font.name = "Courier New"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def parse_docx_to_turns(docx_bytes: bytes) -> List[dict]:
    """
    Parse a DOCX file (exported from TranscribeAlpha) back into transcript turns.

    Expected format per paragraph: "SPEAKER:   Text of what they said..."
    Returns list of dicts with 'speaker' and 'text' keys.

    Automatically skips title page content (Generated Transcript, Case Name, etc.)
    """
    buffer = io.BytesIO(docx_bytes)
    doc = Document(buffer)

    # Title page patterns to skip (case-insensitive)
    title_page_patterns = [
        r'^generated\s+transcript\s*$',
        r'^case\s+name:\s*',
        r'^case\s+number:\s*',
        r'^date:\s*',
        r'^time:\s*',
        r'^location:\s*',
        r'^original\s+file:\s*',
        r'^duration:\s*',
        r'^firm\s*(name|or\s+organization)?\s*:\s*',
    ]
    title_page_regex = re.compile('|'.join(title_page_patterns), re.IGNORECASE)

    turns = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Skip title page content
        if title_page_regex.match(text):
            continue

        # Look for speaker pattern: "SPEAKER:   text" (colon + spaces)
        # Handle various spacing patterns
        match = re.match(r'^([A-Z][A-Z0-9\s\-\.\']*?):\s{1,5}(.+)$', text, re.IGNORECASE)
        if match:
            speaker = match.group(1).strip().upper()
            content = match.group(2).strip()
            if speaker and content:
                # Check if same speaker as previous turn
                is_continuation = False
                if turns and turns[-1]['speaker'] == speaker:
                    is_continuation = True
                turns.append({
                    'speaker': speaker,
                    'text': content,
                    'is_continuation': is_continuation,
                })
        else:
            # No speaker pattern found - treat as continuation of previous turn
            if turns:
                prev = turns[-1]
                prev['text'] = (prev['text'] + ' ' + text).strip()
            else:
                # No previous turn - create a generic one
                turns.append({
                    'speaker': 'SPEAKER',
                    'text': text,
                    'is_continuation': False,
                })

    logger.info("Parsed %d turns from DOCX (skipped title page content)", len(turns))
    return turns


def calculate_line_timestamps_from_words(
    text_line: str,
    all_words: List[WordTimestamp],
    start_offset: int = 0,
) -> Tuple[float, float, int, bool]:
    """
    Calculate accurate start/stop timestamps for a wrapped line using word-level data.

    This function matches words in the text line to their precise timestamps from
    word-level data, eliminating the need for linear interpolation.

    Args:
        text_line: The text content of the line (without speaker prefix)
        all_words: List of WordTimestamp objects for the entire speaker turn
        start_offset: Index in all_words to start searching from (for continuation lines)

    Returns:
        Tuple of (start_seconds, stop_seconds, words_consumed, boundary_missing)
        - start_seconds: Start time of first word in line (seconds, float)
        - stop_seconds: End time of last word in line (seconds, float)
        - words_consumed: Number of words from all_words that were used
        - boundary_missing: True if the first or last word in the line lacks a valid timestamp
    """
    if not all_words or not text_line.strip():
        logger.debug("calculate_line_timestamps: empty words or text_line")
        return (0.0, 0.0, 0, True)

    line_text_clean = text_line.strip().lower()
    if not line_text_clean:
        return (0.0, 0.0, 0, True)

    raw_line_words = line_text_clean.split()

    def normalize_word_for_match(word: str) -> str:
        normalized = word.lower()
        normalized = normalized.replace("’", "'").replace("‘", "'")
        normalized = re.sub(r"[^\w]+", "", normalized)
        return normalized.strip("_")

    line_words = [word for word in raw_line_words if normalize_word_for_match(word)]
    if not line_words:
        return (0.0, 0.0, 0, True)

    # Matching words in order with flexible punctuation handling
    matched_words: List[WordTimestamp] = []
    word_idx = start_offset
    line_idx = 0

    while word_idx < len(all_words) and line_idx < len(line_words):
        line_word = normalize_word_for_match(line_words[line_idx])
        if not line_word:
            line_idx += 1
            continue

        current_word = all_words[word_idx]
        current_word_clean = normalize_word_for_match(current_word.text)

        if current_word_clean == line_word:
            matched_words.append(current_word)
            line_idx += 1
        else:
            if line_word in current_word_clean or current_word_clean in line_word:
                matched_words.append(current_word)
                line_idx += 1

        word_idx += 1

    if not matched_words:
        logger.debug("No words matched for line: %s", text_line)
        return (0.0, 0.0, 0, True)

    first_word = matched_words[0]
    last_word = matched_words[-1]

    start_seconds = (first_word.start or 0.0) / 1000.0
    stop_seconds = (last_word.end or 0.0) / 1000.0
    words_consumed = len(matched_words)

    boundary_missing = False
    if first_word.start is None or last_word.end is None:
        boundary_missing = True

    return (start_seconds, stop_seconds, words_consumed, boundary_missing)


def enforce_min_line_durations(
    line_entries: List[dict],
    audio_duration: float,
    min_duration: float = MIN_LINE_DURATION_SECONDS,
) -> List[dict]:
    if not line_entries or min_duration <= 0:
        return line_entries

    starts: List[float] = []
    ends: List[float] = []
    for entry in line_entries:
        try:
            start_val = float(entry.get("start", 0.0))
        except (TypeError, ValueError):
            start_val = 0.0
        try:
            end_val = float(entry.get("end", start_val))
        except (TypeError, ValueError):
            end_val = start_val
        if end_val < start_val:
            end_val = start_val
        starts.append(start_val)
        ends.append(end_val)

    count = len(line_entries)
    gaps: List[float] = []
    for idx in range(count - 1):
        gap = starts[idx + 1] - ends[idx]
        gaps.append(gap if gap > 0 else 0.0)

    left_gap = [0.0] * count
    right_gap = [0.0] * count

    left_gap[0] = starts[0] if starts[0] > 0 else 0.0
    for idx in range(1, count):
        left_gap[idx] = gaps[idx - 1]
    for idx in range(count - 1):
        right_gap[idx] = gaps[idx]
    if audio_duration and audio_duration > 0:
        right_gap[count - 1] = max(audio_duration - ends[count - 1], 0.0)
    else:
        right_gap[count - 1] = 0.0

    desired_left = [0.0] * count
    desired_right = [0.0] * count

    for idx in range(count):
        duration = ends[idx] - starts[idx]
        if duration >= min_duration:
            continue
        need = min_duration - duration
        half = need / 2.0
        left_take = min(half, left_gap[idx])
        right_take = min(half, right_gap[idx])
        remaining = need - (left_take + right_take)
        if remaining > 0:
            left_cap = max(left_gap[idx] - left_take, 0.0)
            right_cap = max(right_gap[idx] - right_take, 0.0)
            if right_cap > left_cap:
                extra_right = min(remaining, right_cap)
                right_take += extra_right
                remaining -= extra_right
            extra_left = min(remaining, left_cap)
            left_take += extra_left
        desired_left[idx] = left_take
        desired_right[idx] = right_take

    left_alloc = [0.0] * count
    right_alloc = [0.0] * count
    left_alloc[0] = min(desired_left[0], left_gap[0])
    right_alloc[count - 1] = min(desired_right[count - 1], right_gap[count - 1])

    for idx in range(count - 1):
        gap = gaps[idx]
        total = desired_right[idx] + desired_left[idx + 1]
        if total <= 0:
            right_alloc[idx] = 0.0
            left_alloc[idx + 1] = 0.0
        elif total <= gap:
            right_alloc[idx] = desired_right[idx]
            left_alloc[idx + 1] = desired_left[idx + 1]
        else:
            scale = gap / total if total > 0 else 0.0
            right_alloc[idx] = desired_right[idx] * scale
            left_alloc[idx + 1] = desired_left[idx + 1] * scale

    for idx, entry in enumerate(line_entries):
        new_start = starts[idx] - left_alloc[idx]
        new_end = ends[idx] + right_alloc[idx]
        if new_start < 0:
            new_start = 0.0
        if audio_duration and audio_duration > 0 and new_end > audio_duration:
            new_end = audio_duration
        if new_end < new_start:
            new_end = new_start
        entry["start"] = new_start
        entry["end"] = new_end

    return line_entries


def compute_transcript_line_entries(
    transcript_turns: List[TranscriptTurn],
    audio_duration: float,
    lines_per_page: int = 25,
    enforce_min_duration: bool = True,
) -> Tuple[List[dict], int]:
    """
    Build per-line timing/layout data from transcript turns for re-use in XML and editor exports.
    """
    line_entries: List[dict] = []
    page = 1
    line_in_page = 1
    last_pgln = ONCUE_FIRST_PGLN

    for turn_idx, turn in enumerate(transcript_turns):
        start_sec = timestamp_to_seconds(turn.timestamp)
        stop_sec: Optional[float] = None

        if turn.words:
            word_starts = [word.start for word in turn.words if word.start is not None and word.start >= 0]
            word_ends = [word.end for word in turn.words if word.end is not None and word.end >= 0]
            if word_starts and word_ends:
                start_sec = min(word_starts) / 1000.0
                stop_sec = max(word_ends) / 1000.0

        if stop_sec is None:
            if turn_idx < len(transcript_turns) - 1:
                stop_sec = timestamp_to_seconds(transcript_turns[turn_idx + 1].timestamp)
            else:
                stop_sec = audio_duration

        if stop_sec < start_sec:
            stop_sec = start_sec

        speaker_name = turn.speaker.upper()
        text = turn.text.strip()

        # Check if this turn is a continuation of the same speaker
        is_turn_continuation = getattr(turn, 'is_continuation', False)

        if is_turn_continuation:
            # Continuation turn: no speaker prefix, use continuation formatting
            speaker_prefix = ""
            max_first_line_text = MAX_CONTINUATION_WIDTH - CONTINUATION_SPACES
        else:
            # New speaker: include speaker prefix
            speaker_prefix = " " * SPEAKER_PREFIX_SPACES + speaker_name + SPEAKER_COLON
            max_first_line_text = MAX_TOTAL_LINE_WIDTH - len(speaker_prefix)

        wrapped_lines = wrap_text_for_transcript(text, max_first_line_text)
        if not wrapped_lines:
            wrapped_lines = [""]

        max_continuation_text = MAX_CONTINUATION_WIDTH - CONTINUATION_SPACES
        remaining_text = " ".join(wrapped_lines[1:])
        continuation_wrapped = []
        if remaining_text:
            continuation_wrapped = wrap_text_for_transcript(remaining_text, max_continuation_text)

        all_lines = [wrapped_lines[0]] + continuation_wrapped
        total_lines = len(all_lines)

        def interpolate_line_block(block_start: float, block_end: float, count: int) -> List[Tuple[float, float]]:
            if count <= 0:
                return []
            if block_end < block_start:
                block_end = block_start
            duration = block_end - block_start
            step = duration / count if count > 0 else duration
            return [(block_start + step * idx, block_start + step * (idx + 1)) for idx in range(count)]

        line_timings: List[Optional[Tuple[float, float]]] = []
        line_errors: List[bool] = []

        if turn.words:
            word_offset = 0

            for line_text in all_lines:
                line_start, line_stop, words_used, boundary_missing = calculate_line_timestamps_from_words(
                    line_text,
                    turn.words,
                    word_offset,
                )
                if words_used == 0:
                    line_timings.append(None)
                    line_errors.append(True)
                else:
                    line_timings.append((line_start, line_stop))
                    line_errors.append(boundary_missing)
                    word_offset += words_used
        else:
            logger.warning("Turn %d has NO word data, using interpolation", turn_idx)
            line_timings = [None for _ in range(total_lines)]
            line_errors = [True for _ in range(total_lines)]

        if any(timing is None for timing in line_timings):
            filled_timings: List[Tuple[float, float]] = []
            idx = 0
            prev_end = start_sec
            while idx < total_lines:
                timing = line_timings[idx]
                if timing is not None:
                    filled_timings.append(timing)
                    prev_end = timing[1]
                    idx += 1
                    continue

                next_idx = idx
                while next_idx < total_lines and line_timings[next_idx] is None:
                    next_idx += 1
                next_start = line_timings[next_idx][0] if next_idx < total_lines else stop_sec
                block = interpolate_line_block(prev_end, next_start, next_idx - idx)
                filled_timings.extend(block)
                if block:
                    prev_end = block[-1][1]
                idx = next_idx

            line_timings = [timing for timing in filled_timings]

        first_line_start, first_line_stop = line_timings[0] if line_timings else (start_sec, start_sec)
        continuation_timings: List[Tuple[float, float]] = []
        for cont_idx in range(1, total_lines):
            continuation_timings.append(line_timings[cont_idx])

        # First line of turn (with or without speaker prefix depending on continuation status)
        pgln = page * 100 + line_in_page
        last_pgln = pgln

        if is_turn_continuation:
            # Continuation turn: format like a continuation line (no speaker)
            rendered_first_line = " " * CONTINUATION_SPACES + wrapped_lines[0]
        else:
            # New speaker: include speaker prefix
            rendered_first_line = speaker_prefix + wrapped_lines[0]

        line_entries.append(
            {
                "id": f"{turn_idx}-0",
                "turn_index": turn_idx,
                "line_index": 0,
                "speaker": speaker_name,
                "text": wrapped_lines[0],
                "rendered_text": rendered_first_line,
                "start": first_line_start,
                "end": first_line_stop,
                "page": page,
                "line": line_in_page,
                "pgln": pgln,
                "is_continuation": is_turn_continuation,
                "total_lines_in_turn": total_lines,
                "timestamp_error": line_errors[0] if line_errors else False,
            }
        )

        line_in_page += 1
        if line_in_page > lines_per_page:
            page += 1
            line_in_page = 1

        # Continuation lines
        for cont_idx, continuation_text in enumerate(continuation_wrapped):
            line_start, line_stop = continuation_timings[cont_idx]

            pgln = page * 100 + line_in_page
            last_pgln = pgln
            line_entries.append(
                {
                    "id": f"{turn_idx}-{cont_idx + 1}",
                    "turn_index": turn_idx,
                    "line_index": cont_idx + 1,
                    "speaker": speaker_name,
                    "text": continuation_text,
                    "rendered_text": " " * CONTINUATION_SPACES + continuation_text,
                    "start": line_start,
                    "end": line_stop,
                    "page": page,
                    "line": line_in_page,
                    "pgln": pgln,
                    "is_continuation": True,
                    "total_lines_in_turn": total_lines,
                    "timestamp_error": line_errors[cont_idx + 1] if line_errors else False,
                }
            )

            line_in_page += 1
            if line_in_page > lines_per_page:
                page += 1
                line_in_page = 1

    if enforce_min_duration:
        enforce_min_line_durations(line_entries, audio_duration, MIN_LINE_DURATION_SECONDS)

    return line_entries, last_pgln


def generate_oncue_xml(
    transcript_turns: List[TranscriptTurn],
    metadata: dict,
    audio_duration: float,
    lines_per_page: int = 25,
    enforce_min_duration: bool = True,
) -> str:
    """
    Generate OnCue-compatible XML from transcript turns.

    This function breaks long utterances into multiple lines to match the DOCX formatting:
    - First line: 15 spaces + "SPEAKER:" (padded to ~21 chars) + text (max ~48 chars)
    - Continuation lines: 5 spaces + text (max ~57 chars)
    - Total line length: ~71 chars for speaker lines, ~62 chars for continuation lines
    """
    from xml.etree.ElementTree import Element, SubElement, tostring

    root = Element(
        "onCue",
        {
            "xmlns:xsd": "http://www.w3.org/2001/XMLSchema",
            "xmlns:xsi": "http://www.w3.org/2001/XMLSchema-instance",
        },
    )

    media_id = metadata.get("MEDIA_ID") or os.path.splitext(metadata.get("FILE_NAME", "deposition"))[0]
    depo_attrs = {
        "mediaId": str(media_id),
        "linesPerPage": str(lines_per_page),
    }
    if metadata.get("DATE"):
        depo_attrs["date"] = metadata["DATE"]

    deposition = SubElement(root, "deposition", depo_attrs)

    video_attrs = {
        "ID": DEFAULT_VIDEO_ID,
        "filename": metadata.get("FILE_NAME", "audio.mp3"),
        "startTime": "0",
        "stopTime": str(int(round(audio_duration))),
        "firstPGLN": str(ONCUE_FIRST_PGLN),
        "lastPGLN": "0",  # placeholder
        "startTuned": "no",
        "stopTuned": "no",
    }
    depo_video = SubElement(deposition, "depoVideo", video_attrs)

    line_entries, last_pgln = compute_transcript_line_entries(
        transcript_turns,
        audio_duration,
        lines_per_page,
        enforce_min_duration=enforce_min_duration,
    )

    for entry in line_entries:
        SubElement(
            depo_video,
            "depoLine",
            {
                "prefix": "",
                "text": entry["rendered_text"],
                "page": str(entry["page"]),
                "line": str(entry["line"]),
                "pgLN": str(entry["pgln"]),
                "videoID": DEFAULT_VIDEO_ID,
                "videoStart": f"{entry['start']:.2f}",
                "videoStop": f"{entry['end']:.2f}",
                "isEdited": "no",
                "isSynched": "yes",
                "isRedacted": "no",
            },
        )

    depo_video.set("lastPGLN", str(last_pgln))

    xml_bytes = tostring(root, encoding="utf-8", method="xml")
    xml_str = xml_bytes.decode("utf-8")
    xml_str = "".join(xml_str.splitlines())  # single line like sample
    return xml_str
