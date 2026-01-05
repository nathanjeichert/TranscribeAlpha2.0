import os
import io
import json
import time
import re
import tempfile
import logging
import shutil
from typing import List, Optional
import sys

# Python 3.9+ type hint compatibility
if sys.version_info >= (3, 9):
    from typing import Tuple
else:
    from typing import Tuple as typing_Tuple
    Tuple = typing_Tuple

from docx import Document
from docx.shared import Inches, Pt
import ffmpeg
from pydub import AudioSegment
from pydantic import BaseModel

# AssemblyAI integration
try:
    import assemblyai as aai
    ASSEMBLYAI_AVAILABLE = True
except ImportError:
    ASSEMBLYAI_AVAILABLE = False
    logging.getLogger(__name__).warning("AssemblyAI SDK not installed. Run: pip install assemblyai")

# Configure both ffmpeg libraries to find ffmpeg
import subprocess
import platform

def find_executable_path(executable_name: str) -> Optional[str]:
    """Cross-platform executable finder"""
    # First try shutil.which (works on all platforms)
    path = shutil.which(executable_name)
    if path:
        return path
    
    # Platform-specific fallback commands
    system = platform.system().lower()
    if system == "windows":
        try:
            result = subprocess.run(['where', executable_name], capture_output=True, text=True, shell=True)
            if result.returncode == 0:
                return result.stdout.strip().split('\n')[0]
        except Exception:
            pass
    else:
        try:
            result = subprocess.run(['which', executable_name], capture_output=True, text=True)
            if result.returncode == 0:
                return result.stdout.strip()
        except Exception:
            pass
    
    return None

def get_ffprobe_path(ffmpeg_path: str) -> Optional[str]:
    """Get ffprobe path from ffmpeg path"""
    if not ffmpeg_path:
        return None
    
    # Get directory and base name
    ffmpeg_dir = os.path.dirname(ffmpeg_path)
    ffmpeg_name = os.path.basename(ffmpeg_path)
    
    # Replace ffmpeg with ffprobe, keeping the same extension
    if ffmpeg_name.endswith('.exe'):
        ffprobe_name = ffmpeg_name.replace('ffmpeg.exe', 'ffprobe.exe')
    else:
        ffprobe_name = ffmpeg_name.replace('ffmpeg', 'ffprobe')
    
    ffprobe_path = os.path.join(ffmpeg_dir, ffprobe_name)
    
    # Check if it exists
    if os.path.exists(ffprobe_path):
        return ffprobe_path
    
    # Try finding ffprobe independently
    return find_executable_path('ffprobe')

# Find ffmpeg and ffprobe
ffmpeg_executable_path = find_executable_path('ffmpeg')
ffprobe_executable_path = None

if ffmpeg_executable_path:
    ffprobe_executable_path = get_ffprobe_path(ffmpeg_executable_path)
    # Configure pydub
    AudioSegment.converter = ffmpeg_executable_path
    AudioSegment.ffmpeg = ffmpeg_executable_path
    if ffprobe_executable_path:
        AudioSegment.ffprobe = ffprobe_executable_path

SUPPORTED_VIDEO_TYPES = ["mp4", "mov", "avi", "mkv"]
SUPPORTED_AUDIO_TYPES = ["mp3", "wav", "m4a", "flac", "ogg", "aac", "aiff"]

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

ASSEMBLYAI_API_KEY = os.getenv("ASSEMBLYAI_API_KEY")
if ASSEMBLYAI_API_KEY and ASSEMBLYAI_AVAILABLE:
    aai.settings.api_key = ASSEMBLYAI_API_KEY
    logger.info("AssemblyAI client initialized successfully")
elif ASSEMBLYAI_AVAILABLE:
    logger.warning("ASSEMBLYAI_API_KEY environment variable not set")

class WordTimestamp(BaseModel):
    """Represents a single word with precise timing information"""
    text: str
    start: float  # Start time in milliseconds
    end: float    # End time in milliseconds
    confidence: Optional[float] = None
    speaker: Optional[str] = None

class TranscriptTurn(BaseModel):
    speaker: str
    text: str
    timestamp: Optional[str] = None
    words: Optional[List[WordTimestamp]] = None  # Word-level timestamps for accurate line timing
    is_continuation: bool = False  # True if same speaker as previous turn (no speaker label needed)


def mark_continuation_turns(turns: List[TranscriptTurn]) -> List[TranscriptTurn]:
    """
    Mark turns as continuations when the same speaker has consecutive turns.

    The first turn of each speaker block gets is_continuation=False (shows speaker label).
    Subsequent turns with the same speaker get is_continuation=True (no speaker label).
    """
    if not turns:
        return turns

    prev_speaker = None
    for turn in turns:
        normalized_speaker = turn.speaker.strip().upper()
        if prev_speaker is not None and normalized_speaker == prev_speaker:
            turn.is_continuation = True
        else:
            turn.is_continuation = False
        prev_speaker = normalized_speaker

    return turns


def convert_video_to_audio(input_path: str, output_path: str, format: str = "mp3") -> Optional[str]:
    try:
        logger.info("Converting %s to %s", input_path, output_path)

        if ffmpeg_executable_path:
            cmd = [
                ffmpeg_executable_path,
                '-i', input_path,
                '-acodec', 'libmp3lame',
                '-y',
                output_path
            ]
            logger.debug("Running command: %s", ' '.join(cmd))
            result = subprocess.run(cmd, capture_output=True, text=True)
            if result.returncode == 0:
                logger.info("Successfully converted to %s", output_path)
                time.sleep(0.5)
                return output_path
            else:
                logger.error("ffmpeg failed with return code %d: %s", result.returncode, result.stderr)
                return None
        else:
            logger.debug("Using ffmpeg-python library as fallback")
            ffmpeg.input(input_path).output(output_path, format=format, acodec='libmp3lame').overwrite_output().run(quiet=True)
            return output_path
    except Exception as e:
        logger.error("Unexpected error in convert_video_to_audio: %s", e)
        return None


def get_audio_mime_type(ext: str) -> Optional[str]:
    mime_map = {
        "mp3": "audio/mpeg",
        "wav": "audio/wav",
        "aiff": "audio/aiff",
        "aac": "audio/aac",
        "ogg": "audio/ogg",
        "flac": "audio/flac",
    }
    return mime_map.get(ext.lower())


def transcribe_with_assemblyai(
    audio_path: str,
    speaker_name_list: Optional[List[str]] = None,
    include_timestamps: bool = True
) -> Optional[List[TranscriptTurn]]:
    """
    Transcribe audio using AssemblyAI with speaker diarization and word-level timestamps.

    Args:
        audio_path: Path to audio file (local file path)
        speaker_name_list: Optional list of speaker names to map to AssemblyAI's labels
        include_timestamps: Whether to include timestamps in output

    Returns:
        List of TranscriptTurn objects with word-level timing data, or None on failure

    Note:
        AssemblyAI labels speakers as "A", "B", "C", etc. This function maps them to
        provided speaker names or generates generic identifiers like "SPEAKER A".
    """
    if not ASSEMBLYAI_AVAILABLE:
        logger.error("AssemblyAI SDK not available")
        return None

    if not ASSEMBLYAI_API_KEY:
        logger.error("ASSEMBLYAI_API_KEY not configured")
        return None

    try:
        # Configure transcription with speaker diarization
        raw_config = aai.RawTranscriptionConfig(
            language_model="slam_1",
            acoustic_model="slam_1",
        )

        config = aai.TranscriptionConfig(
            speaker_labels=True,
            speakers_expected=len(speaker_name_list) if speaker_name_list else None,
            raw_transcription_config=raw_config,
        )

        logger.info(f"Starting AssemblyAI transcription for: {audio_path}")
        logger.info(
            "Speaker diarization enabled, expected speakers: %s",
            len(speaker_name_list) if speaker_name_list else "auto-detect",
        )

        # Transcribe audio file
        transcriber = aai.Transcriber()
        transcript = transcriber.transcribe(audio_path, config=config)

        # Check for errors
        if transcript.status == aai.TranscriptStatus.error:
            logger.error("AssemblyAI transcription failed: %s", transcript.error)
            return None

        logger.info("AssemblyAI transcription completed successfully")
        logger.info(
            "Found %s speaker turns",
            len(transcript.utterances) if transcript.utterances else 0,
        )

        # Convert AssemblyAI utterances to TranscriptTurn format
        turns: List[TranscriptTurn] = []

        for utterance in transcript.utterances or []:
            # Map AssemblyAI speaker labels (A, B, C...) to provided names
            speaker_label = utterance.speaker  # e.g., "A", "B", "C"

            if speaker_name_list and speaker_label:
                try:
                    speaker_idx = ord(speaker_label) - ord("A")
                    if 0 <= speaker_idx < len(speaker_name_list):
                        speaker_name = speaker_name_list[speaker_idx]
                    else:
                        speaker_name = f"SPEAKER {speaker_label}"
                except (TypeError, ValueError):
                    speaker_name = f"SPEAKER {speaker_label}"
            else:
                speaker_name = f"SPEAKER {speaker_label}" if speaker_label else "SPEAKER 1"

            # Convert timestamp from milliseconds to [MM:SS] format for consistency
            timestamp_str = None
            if include_timestamps and getattr(utterance, "start", None) is not None:
                start_ms = utterance.start
                start_seconds = start_ms / 1000.0
                minutes = int(start_seconds // 60)
                seconds = int(start_seconds % 60)
                timestamp_str = f"[{minutes:02d}:{seconds:02d}]"

            # Extract word-level timestamps from utterance
            word_timestamps: List[WordTimestamp] = []
            if hasattr(utterance, "words") and utterance.words:
                for word in utterance.words:
                    word_timestamps.append(
                        WordTimestamp(
                            text=word.text,
                            start=float(word.start),
                            end=float(word.end),
                            confidence=float(word.confidence)
                            if hasattr(word, "confidence") and word.confidence is not None
                            else None,
                            speaker=speaker_name,
                        )
                    )

            turns.append(
                TranscriptTurn(
                    speaker=speaker_name,
                    text=utterance.text,
                    timestamp=timestamp_str,
                    words=word_timestamps if word_timestamps else None,
                )
            )

        logger.info("Converted %s utterances to TranscriptTurn format", len(turns))
        # Mark continuation turns (same speaker as previous)
        turns = mark_continuation_turns(turns)
        return turns

    except Exception as e:
        logger.error("AssemblyAI transcription error: %s", str(e))
        import traceback

        logger.error(traceback.format_exc())
        return None


def calculate_line_timestamps_from_words(
    text_line: str,
    all_words: List[WordTimestamp],
    start_offset: int = 0
) -> Tuple[float, float, int]:
    """
    Calculate accurate start/stop timestamps for a wrapped line using word-level data.

    This function matches words in the text line to their precise timestamps from
    word-level data, eliminating the need for linear interpolation.

    Args:
        text_line: The text content of the line (without speaker prefix)
        all_words: List of WordTimestamp objects for the entire speaker turn
        start_offset: Index in all_words to start searching from (for continuation lines)

    Returns:
        Tuple of (start_seconds, stop_seconds, words_consumed)
        - start_seconds: Start time of first word in line (seconds, float)
        - stop_seconds: End time of last word in line (seconds, float)
        - words_consumed: Number of words from all_words that were used
    """
    if not all_words or not text_line.strip():
        logger.debug("calculate_line_timestamps: empty words or text_line")
        return (0.0, 0.0, 0)

    line_text_clean = text_line.strip().lower()
    if not line_text_clean:
        return (0.0, 0.0, 0)

    line_words = line_text_clean.split()
    if not line_words:
        return (0.0, 0.0, 0)

    matched_word_indices: List[int] = []
    word_search_idx = start_offset
    line_word_idx = 0

    def normalize_word_for_match(word: str) -> str:
        normalized = word.lower()
        normalized = normalized.replace("’", "'").replace("‘", "'")
        normalized = re.sub(r"[^\w]+", "", normalized)
        return normalized.strip("_")

    while line_word_idx < len(line_words) and word_search_idx < len(all_words):
        word_obj = all_words[word_search_idx]
        word_clean = normalize_word_for_match(word_obj.text)
        line_word_clean = normalize_word_for_match(line_words[line_word_idx])

        if word_clean == line_word_clean:
            matched_word_indices.append(word_search_idx)
            line_word_idx += 1
            word_search_idx += 1
        else:
            word_search_idx += 1

    if not matched_word_indices:
        logger.warning("No word matches found for line: '%s...' (first word in array: '%s')",
                      text_line[:50], all_words[start_offset].text if start_offset < len(all_words) else 'N/A')
        if start_offset < len(all_words):
            word = all_words[start_offset]
            return (word.start / 1000.0, word.end / 1000.0, 1)
        return (0.0, 0.0, 0)

    first_word = all_words[matched_word_indices[0]]
    last_word = all_words[matched_word_indices[-1]]

    start_seconds = first_word.start / 1000.0
    stop_seconds = last_word.end / 1000.0
    words_consumed = len(matched_word_indices)

    logger.debug("Line timestamps: %.2f-%.2f (%d words matched) for '%s...'",
                start_seconds, stop_seconds, words_consumed, text_line[:30])

    return (start_seconds, stop_seconds, words_consumed)


def replace_placeholder_text(element, placeholder: str, replacement: str) -> None:
    if hasattr(element, 'paragraphs'):
        for p in element.paragraphs:
            replace_placeholder_text(p, placeholder, replacement)
    if hasattr(element, 'runs'):
        if placeholder in element.text:
            inline = element.runs
            for i in range(len(inline)):
                if placeholder in inline[i].text:
                    text = inline[i].text.replace(placeholder, replacement)
                    inline[i].text = text
    if hasattr(element, 'tables'):
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_placeholder_text(cell, placeholder, replacement)


def create_docx(title_data: dict, transcript_turns: List[TranscriptTurn]) -> bytes:
    """
    Create a DOCX transcript with simple one-paragraph-per-turn formatting.

    Word will automatically wrap lines based on the first-line indent (1.0").
    The XML generation must match Word's natural line wrapping behavior.
    Timestamps are currently excluded to keep pagination aligned with OnCue output.
    """
    doc = Document("transcript_template.docx")
    for key, value in title_data.items():
        placeholder = f"{{{{{key}}}}}"
        replace_placeholder_text(doc, placeholder, str(value) if value else "")

    body_placeholder = "{{TRANSCRIPT_BODY}}"
    placeholder_paragraph = None
    for p in doc.paragraphs:
        if body_placeholder in p.text:
            placeholder_paragraph = p
            break

    if placeholder_paragraph:
        p_element = placeholder_paragraph._element
        p_element.getparent().remove(p_element)
        for turn in transcript_turns:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.first_line_indent = Inches(1.0)  # Standard legal transcript indent
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.widow_control = False

            # Only show speaker label if not a continuation turn
            if not turn.is_continuation:
                speaker_run = p.add_run(f"{turn.speaker.upper()}:   ")
                speaker_run.font.name = "Courier New"
            text_run = p.add_run(turn.text)
            text_run.font.name = "Courier New"
    else:
        for turn in transcript_turns:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.first_line_indent = Inches(1.0)
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.widow_control = False

            # Only show speaker label if not a continuation turn
            if not turn.is_continuation:
                speaker_run = p.add_run(f"{turn.speaker.upper()}:   ")
                speaker_run.font.name = "Courier New"
            text_run = p.add_run(turn.text)
            text_run.font.name = "Courier New"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def parse_docx_to_turns(docx_bytes: bytes) -> List[dict]:
    """
    Parse a DOCX file (exported from TranscribeAlpha) back into transcript turns.

    Expected format per paragraph: "SPEAKER:   Text of what they said..."
    Returns list of dicts with 'speaker' and 'text' keys.

    Automatically skips title page content (Generated Transcript, Case Name, etc.)
    """
    buffer = io.BytesIO(docx_bytes)
    doc = Document(buffer)

    # Title page patterns to skip (case-insensitive)
    title_page_patterns = [
        r'^generated\s+transcript\s*$',
        r'^case\s+name:\s*',
        r'^case\s+number:\s*',
        r'^date:\s*',
        r'^time:\s*',
        r'^location:\s*',
        r'^original\s+file:\s*',
        r'^duration:\s*',
        r'^firm\s*(name|or\s+organization)?\s*:\s*',
    ]
    title_page_regex = re.compile('|'.join(title_page_patterns), re.IGNORECASE)

    turns = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Skip title page content
        if title_page_regex.match(text):
            continue

        # Look for speaker pattern: "SPEAKER:   text" (colon + spaces)
        # Handle various spacing patterns
        match = re.match(r'^([A-Z][A-Z0-9\s\-\.\']*?):\s{1,5}(.+)$', text, re.IGNORECASE)
        if match:
            speaker = match.group(1).strip().upper()
            content = match.group(2).strip()
            if speaker and content:
                # Check if same speaker as previous turn
                is_continuation = (turns and turns[-1]['speaker'].upper() == speaker)
                turns.append({
                    'speaker': speaker,
                    'text': content,
                    'is_continuation': is_continuation,
                })
        else:
            # No speaker pattern - this is a continuation of the previous speaker
            if turns and not text.startswith('['):  # Avoid timestamps
                # Create as separate continuation turn (inherits speaker from previous)
                turns.append({
                    'speaker': turns[-1]['speaker'],
                    'text': text,
                    'is_continuation': True,
                })
            elif text and not text.startswith('['):
                turns.append({
                    'speaker': 'UNKNOWN',
                    'text': text,
                    'is_continuation': False,
                })

    logger.info(f"Parsed {len(turns)} turns from DOCX (skipped title page content)")
    return turns


# Helper to get media duration with ffprobe, avoids pydub's internal lookup issues

def get_media_duration(file_path: str) -> Optional[float]:
    """Return the duration of an audio/video file in **seconds** using ffprobe.
    
    Uses the cross-platform ffprobe path detection.
    """
    if not ffprobe_executable_path:
        return None  # ffprobe not available

    try:
        cmd = [
            ffprobe_executable_path,
            "-i",
            file_path,
            "-v",
            "error",
            "-show_entries",
            "format=duration",
            "-of",
            "default=noprint_wrappers=1:nokey=1",
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        duration_str = result.stdout.strip()
        if duration_str:
            return float(duration_str)
    except Exception as e:
        logger.debug("ffprobe duration extraction failed: %s", e)
    return None

def process_transcription(
    file_bytes: bytes,
    filename: str,
    speaker_names: Optional[List[str]],
    title_data: dict,
):
    with tempfile.TemporaryDirectory() as temp_dir:
        input_path = os.path.join(temp_dir, filename)
        with open(input_path, "wb") as f:
            f.write(file_bytes)

        ext = filename.split('.')[-1].lower()
        audio_path = None
        if ext in SUPPORTED_VIDEO_TYPES:
            output_audio_filename = f"{os.path.splitext(filename)[0]}.mp3"
            output_path = os.path.join(temp_dir, output_audio_filename)
            audio_path = convert_video_to_audio(input_path, output_path)
            ext = "mp3"
        elif ext in SUPPORTED_AUDIO_TYPES:
            audio_path = input_path
        else:
            raise ValueError("Unsupported file type")

        mime_type = get_audio_mime_type(ext)

        # ------------------------------------------------------------------
        # Retrieve media duration – prefer direct ffprobe for robustness
        # ------------------------------------------------------------------
        duration_seconds = get_media_duration(audio_path)

        if duration_seconds is None:
            # Fallback to pydub if ffprobe failed for some reason
            audio_segment = None
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    audio_segment = AudioSegment.from_file(audio_path)
                    break
                except (PermissionError, FileNotFoundError) as e:
                    if attempt < max_retries - 1:
                        logger.warning("Attempt %d failed to load audio file: %s. Retrying...", attempt + 1, e)
                        time.sleep(1)
                    else:
                        raise e
            if audio_segment is None:
                raise RuntimeError("Failed to load audio file after multiple attempts")

            duration_seconds = len(audio_segment) / 1000.0

        # ------------------------------------------------------------------
        # Format and store duration for title data
        # ------------------------------------------------------------------
        hours, rem = divmod(duration_seconds, 3600)
        minutes, seconds = divmod(rem, 60)
        file_duration_str = "{:0>2}:{:0>2}:{:0>2}".format(int(hours), int(minutes), int(round(seconds)))
        title_data["FILE_DURATION"] = file_duration_str

        # ------------------------------------------------------------------
        # Proceed with upload & transcription (AssemblyAI only)
        # ------------------------------------------------------------------
        logger.info("Using AssemblyAI transcription engine")

        if not ASSEMBLYAI_AVAILABLE:
            raise RuntimeError("AssemblyAI SDK not installed. Run: pip install assemblyai")

        if not ASSEMBLYAI_API_KEY:
            raise RuntimeError("ASSEMBLYAI_API_KEY environment variable not set")

        turns = transcribe_with_assemblyai(audio_path, speaker_names)

        if not turns:
            raise RuntimeError("AssemblyAI transcription failed")

        docx_bytes = create_docx(title_data, turns)

        return turns, docx_bytes, duration_seconds
def timestamp_to_seconds(timestamp: Optional[str]) -> float:
    """Convert timestamp like '[MM:SS]' or 'MM:SS' to seconds."""
    if not timestamp:
        return 0.0
    ts = timestamp.strip('[]').strip()
    parts = ts.split(':')
    try:
        if len(parts) == 3:
            h, m, s = map(float, parts)
            return h * 3600 + m * 60 + s
        elif len(parts) == 2:
            m, s = map(float, parts)
            return m * 60 + s
        else:
            return float(ts)
    except ValueError:
        return 0.0


def seconds_to_timestamp(seconds: float) -> str:
    """Convert seconds float to OnCue-style [MM:SS] or [HH:MM:SS] timestamp."""
    if seconds < 0:
        seconds = 0.0
    total_seconds = int(round(seconds))
    hours, remainder = divmod(total_seconds, 3600)
    minutes, secs = divmod(remainder, 60)
    if hours > 0:
        return f"[{hours:02d}:{minutes:02d}:{secs:02d}]"
    return f"[{minutes:02d}:{secs:02d}]"


def wrap_text_for_transcript(text: str, max_width: int) -> List[str]:
    """
    Wrap text to fit within max_width characters, preserving word boundaries.

    Args:
        text: The text to wrap
        max_width: Maximum characters per line of text content

    Returns:
        List of wrapped lines
    """
    if not text:
        return [""]

    if max_width <= 0:
        return [text]

    words = text.split()
    lines = []
    current_line = []
    current_length = 0

    for word in words:
        word_length = len(word)
        # +1 for space before word (except first word)
        space_needed = word_length + (1 if current_line else 0)

        if current_length + space_needed <= max_width:
            current_line.append(word)
            current_length += space_needed
        else:
            # Start new line
            if current_line:
                lines.append(" ".join(current_line))
            current_line = [word]
            current_length = word_length

    # Add remaining words
    if current_line:
        lines.append(" ".join(current_line))

    return lines if lines else [""]


# Shared layout constants used for XML generation and editor exports
SPEAKER_PREFIX_SPACES = 10  # Leading spaces before speaker name in XML (visual simulation)
CONTINUATION_SPACES = 0     # Leading spaces for continuation lines in XML (visual simulation)
SPEAKER_COLON = ":   "      # Colon and spaces after speaker name (total 4 chars)
MAX_TOTAL_LINE_WIDTH = 64   # Maximum total characters per XML line for speaker lines
MAX_CONTINUATION_WIDTH = 64 # Maximum total characters per XML line for continuation lines

# OnCue XML constants
ONCUE_FIRST_PGLN = 101      # First page-line number (page 1, line 1 = 101)
DEFAULT_VIDEO_ID = "1"      # Default video ID for single-video transcripts


def compute_transcript_line_entries(
    transcript_turns: List[TranscriptTurn],
    audio_duration: float,
    lines_per_page: int = 25,
) -> Tuple[List[dict], int]:
    """
    Build per-line timing/layout data from transcript turns for re-use in XML and editor exports.
    """
    line_entries: List[dict] = []
    page = 1
    line_in_page = 1
    last_pgln = ONCUE_FIRST_PGLN

    for turn_idx, turn in enumerate(transcript_turns):
        start_sec = timestamp_to_seconds(turn.timestamp)
        stop_sec: Optional[float] = None

        if turn.words:
            word_starts = [word.start for word in turn.words if word.start is not None]
            word_ends = [word.end for word in turn.words if word.end is not None]
            if word_starts and word_ends:
                start_sec = min(word_starts) / 1000.0
                stop_sec = max(word_ends) / 1000.0

        if stop_sec is None:
            if turn_idx < len(transcript_turns) - 1:
                stop_sec = timestamp_to_seconds(transcript_turns[turn_idx + 1].timestamp)
            else:
                stop_sec = audio_duration

        if stop_sec < start_sec:
            stop_sec = start_sec

        speaker_name = turn.speaker.upper()
        text = turn.text.strip()

        # Check if this turn is a continuation of the same speaker
        is_turn_continuation = getattr(turn, 'is_continuation', False)

        if is_turn_continuation:
            # Continuation turn: no speaker prefix, use continuation formatting
            speaker_prefix = ""
            max_first_line_text = MAX_CONTINUATION_WIDTH - CONTINUATION_SPACES
        else:
            # New speaker: include speaker prefix
            speaker_prefix = " " * SPEAKER_PREFIX_SPACES + speaker_name + SPEAKER_COLON
            max_first_line_text = MAX_TOTAL_LINE_WIDTH - len(speaker_prefix)

        wrapped_lines = wrap_text_for_transcript(text, max_first_line_text)
        if not wrapped_lines:
            wrapped_lines = [""]

        max_continuation_text = MAX_CONTINUATION_WIDTH - CONTINUATION_SPACES
        remaining_text = " ".join(wrapped_lines[1:])
        continuation_wrapped = []
        if remaining_text:
            continuation_wrapped = wrap_text_for_transcript(remaining_text, max_continuation_text)

        total_lines = 1 + len(continuation_wrapped)

        if turn.words:
            logger.info("Turn %d has %d words, first word: %s (start=%.1f ms)",
                       turn_idx, len(turn.words), turn.words[0].text, turn.words[0].start)
            word_offset = 0
            use_word_data = True

            first_line_start, first_line_stop, words_used = calculate_line_timestamps_from_words(
                wrapped_lines[0],
                turn.words,
                word_offset,
            )
            if words_used == 0:
                logger.warning("Turn %d: word matching failed, falling back to interpolation", turn_idx)
                use_word_data = False
            else:
                logger.info("Turn %d first line: %.2f-%.2f seconds (%d words)",
                           turn_idx, first_line_start, first_line_stop, words_used)
                word_offset += words_used

            continuation_timings: List[Tuple[float, float]] = []
            if use_word_data:
                for cont_line in continuation_wrapped:
                    line_start, line_stop, words_used = calculate_line_timestamps_from_words(
                        cont_line,
                        turn.words,
                        word_offset,
                    )
                    if words_used == 0:
                        use_word_data = False
                        break
                    continuation_timings.append((line_start, line_stop))
                    word_offset += words_used

            if not use_word_data:
                turn_duration = stop_sec - start_sec
                time_per_line = turn_duration / total_lines if total_lines > 0 else turn_duration
                first_line_start = start_sec
                first_line_stop = start_sec + time_per_line
                continuation_timings = []
                for cont_idx in range(len(continuation_wrapped)):
                    line_start = start_sec + time_per_line * (cont_idx + 1)
                    line_stop = start_sec + time_per_line * (cont_idx + 2)
                    if cont_idx == len(continuation_wrapped) - 1:
                        line_stop = stop_sec
                    continuation_timings.append((line_start, line_stop))
        else:
            logger.warning("Turn %d has NO word data, using interpolation", turn_idx)
            turn_duration = stop_sec - start_sec
            time_per_line = turn_duration / total_lines if total_lines > 0 else turn_duration

            first_line_start = start_sec
            first_line_stop = start_sec + time_per_line

            continuation_timings = []
            for cont_idx in range(len(continuation_wrapped)):
                line_start = start_sec + time_per_line * (cont_idx + 1)
                line_stop = start_sec + time_per_line * (cont_idx + 2)
                if cont_idx == len(continuation_wrapped) - 1:
                    line_stop = stop_sec
                continuation_timings.append((line_start, line_stop))

        # First line of turn (with or without speaker prefix depending on continuation status)
        pgln = page * 100 + line_in_page
        last_pgln = pgln

        if is_turn_continuation:
            # Continuation turn: format like a continuation line (no speaker)
            rendered_first_line = " " * CONTINUATION_SPACES + wrapped_lines[0]
        else:
            # New speaker: include speaker prefix
            rendered_first_line = speaker_prefix + wrapped_lines[0]

        line_entries.append(
            {
                "id": f"{turn_idx}-0",
                "turn_index": turn_idx,
                "line_index": 0,
                "speaker": speaker_name if not is_turn_continuation else "",
                "text": wrapped_lines[0],
                "rendered_text": rendered_first_line,
                "start": first_line_start,
                "end": first_line_stop,
                "page": page,
                "line": line_in_page,
                "pgln": pgln,
                "is_continuation": is_turn_continuation,
                "total_lines_in_turn": total_lines,
            }
        )

        line_in_page += 1
        if line_in_page > lines_per_page:
            page += 1
            line_in_page = 1

        # Continuation lines
        for cont_idx, continuation_text in enumerate(continuation_wrapped):
            line_start, line_stop = continuation_timings[cont_idx]

            pgln = page * 100 + line_in_page
            last_pgln = pgln
            line_entries.append(
                {
                    "id": f"{turn_idx}-{cont_idx + 1}",
                    "turn_index": turn_idx,
                    "line_index": cont_idx + 1,
                    "speaker": speaker_name,
                    "text": continuation_text,
                    "rendered_text": " " * CONTINUATION_SPACES + continuation_text,
                    "start": line_start,
                    "end": line_stop,
                    "page": page,
                    "line": line_in_page,
                    "pgln": pgln,
                    "is_continuation": True,
                    "total_lines_in_turn": total_lines,
                }
            )

            line_in_page += 1
            if line_in_page > lines_per_page:
                page += 1
                line_in_page = 1

    return line_entries, last_pgln


def generate_oncue_xml(transcript_turns: List[TranscriptTurn], metadata: dict, audio_duration: float, lines_per_page: int = 25) -> str:
    """
    Generate OnCue-compatible XML from transcript turns.

    This function breaks long utterances into multiple lines to match the DOCX formatting:
    - First line: 15 spaces + "SPEAKER:" (padded to ~21 chars) + text (max ~48 chars)
    - Continuation lines: 5 spaces + text (max ~57 chars)
    - Total line length: ~71 chars for speaker lines, ~62 chars for continuation lines
    """
    from xml.etree.ElementTree import Element, SubElement, tostring

    root = Element(
        "onCue",
        {
            "xmlns:xsd": "http://www.w3.org/2001/XMLSchema",
            "xmlns:xsi": "http://www.w3.org/2001/XMLSchema-instance",
        },
    )

    media_id = metadata.get("MEDIA_ID") or os.path.splitext(metadata.get("FILE_NAME", "deposition"))[0]
    depo_attrs = {
        "mediaId": str(media_id),
        "linesPerPage": str(lines_per_page),
    }
    if metadata.get("DATE"):
        depo_attrs["date"] = metadata["DATE"]

    deposition = SubElement(root, "deposition", depo_attrs)

    video_attrs = {
        "ID": DEFAULT_VIDEO_ID,
        "filename": metadata.get("FILE_NAME", "audio.mp3"),
        "startTime": "0",
        "stopTime": str(int(round(audio_duration))),
        "firstPGLN": str(ONCUE_FIRST_PGLN),
        "lastPGLN": "0",  # placeholder
        "startTuned": "no",
        "stopTuned": "no",
    }
    depo_video = SubElement(deposition, "depoVideo", video_attrs)

    line_entries, last_pgln = compute_transcript_line_entries(
        transcript_turns,
        audio_duration,
        lines_per_page,
    )

    for entry in line_entries:
        SubElement(
            depo_video,
            "depoLine",
            {
                "prefix": "",
                "text": entry["rendered_text"],
                "page": str(entry["page"]),
                "line": str(entry["line"]),
                "pgLN": str(entry["pgln"]),
                "videoID": DEFAULT_VIDEO_ID,
                "videoStart": f"{entry['start']:.2f}",
                "videoStop": f"{entry['end']:.2f}",
                "isEdited": "no",
                "isSynched": "yes",
                "isRedacted": "no",
            },
        )

    depo_video.set("lastPGLN", str(last_pgln))

    xml_bytes = tostring(root, encoding="utf-8", method="xml")
    xml_str = xml_bytes.decode("utf-8")
    xml_str = "".join(xml_str.splitlines())  # single line like sample
    return xml_str
