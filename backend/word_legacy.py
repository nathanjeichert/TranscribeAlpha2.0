"""
Deprecated Word/DOCX helpers.

This module keeps legacy Word logic isolated from the active PDF pipeline.
New transcript exports should use PDF generation in transcript_formatting.py.
"""

import io
import logging
import os
import re
from typing import List, Optional

from docx import Document
from docx.shared import Inches, Pt

try:
    from .models import TranscriptTurn
except ImportError:
    try:
        from models import TranscriptTurn
    except ImportError:
        import models
        TranscriptTurn = models.TranscriptTurn

logger = logging.getLogger(__name__)


def replace_placeholder_text(element, placeholder: str, replacement: str) -> None:
    if hasattr(element, "paragraphs"):
        for paragraph in element.paragraphs:
            replace_placeholder_text(paragraph, placeholder, replacement)
    if hasattr(element, "runs"):
        if placeholder in element.text:
            inline = element.runs
            for idx in range(len(inline)):
                if placeholder in inline[idx].text:
                    inline[idx].text = inline[idx].text.replace(placeholder, replacement)
    if hasattr(element, "tables"):
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_placeholder_text(cell, placeholder, replacement)


def _resolve_docx_template_path() -> Optional[str]:
    candidates = [
        os.path.join(os.path.dirname(__file__), "..", "transcript_template.docx"),
        os.path.join(os.getcwd(), "transcript_template.docx"),
    ]
    for path in candidates:
        if path and os.path.exists(path):
            return path
    return None


def _resolve_clip_template_path() -> Optional[str]:
    candidates = [
        os.path.join(os.path.dirname(__file__), "..", "clip_template.docx"),
        os.path.join(os.getcwd(), "clip_template.docx"),
    ]
    for path in candidates:
        if path and os.path.exists(path):
            return path
    return None


def create_docx(title_data: dict, transcript_turns: List[TranscriptTurn]) -> bytes:
    """
    Deprecated DOCX export helper retained for legacy workflows only.
    """
    template_path = _resolve_docx_template_path()
    if template_path:
        doc = Document(template_path)
    else:
        logger.warning("DOCX template not found; falling back to a blank document.")
        doc = Document()

    for key, value in title_data.items():
        placeholder = f"{{{{{key}}}}}"
        replace_placeholder_text(doc, placeholder, str(value) if value else "")

    body_placeholder = "{{TRANSCRIPT_BODY}}"
    placeholder_paragraph = None
    for paragraph in doc.paragraphs:
        if body_placeholder in paragraph.text:
            placeholder_paragraph = paragraph
            break

    if placeholder_paragraph:
        paragraph_element = placeholder_paragraph._element
        paragraph_element.getparent().remove(paragraph_element)
        for turn in transcript_turns:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.first_line_indent = Inches(1.0)
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.widow_control = False

            if not turn.is_continuation:
                speaker_run = p.add_run(f"{turn.speaker.upper()}:   ")
                speaker_run.font.name = "Courier New"
            text_run = p.add_run(turn.text)
            text_run.font.name = "Courier New"
    else:
        for turn in transcript_turns:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.first_line_indent = Inches(1.0)
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.widow_control = False

            if not turn.is_continuation:
                speaker_run = p.add_run(f"{turn.speaker.upper()}:   ")
                speaker_run.font.name = "Courier New"
            text_run = p.add_run(turn.text)
            text_run.font.name = "Courier New"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def create_clip_docx(title_data: dict, transcript_turns: List[TranscriptTurn], clip_title: str) -> bytes:
    """
    Deprecated clip DOCX export helper retained for legacy workflows only.
    """
    template_path = _resolve_clip_template_path()
    if not template_path:
        logger.warning("Clip template not found; falling back to standard template.")
        template_path = _resolve_docx_template_path()

    if template_path:
        doc = Document(template_path)
    else:
        logger.warning("No DOCX template found; using blank document.")
        doc = Document()

    clip_title_data = dict(title_data)
    clip_title_data["CLIP_TITLE"] = clip_title

    for key, value in clip_title_data.items():
        placeholder = f"{{{{{key}}}}}"
        replace_placeholder_text(doc, placeholder, str(value) if value else "")

    body_placeholder = "{{TRANSCRIPT_BODY}}"
    placeholder_paragraph = None
    for paragraph in doc.paragraphs:
        if body_placeholder in paragraph.text:
            placeholder_paragraph = paragraph
            break

    if placeholder_paragraph:
        paragraph_element = placeholder_paragraph._element
        paragraph_element.getparent().remove(paragraph_element)
        for turn in transcript_turns:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.first_line_indent = Inches(1.0)
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.widow_control = False

            if not turn.is_continuation:
                speaker_run = p.add_run(f"{turn.speaker.upper()}:   ")
                speaker_run.font.name = "Courier New"
            text_run = p.add_run(turn.text)
            text_run.font.name = "Courier New"
    else:
        for turn in transcript_turns:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.first_line_indent = Inches(1.0)
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.widow_control = False

            if not turn.is_continuation:
                speaker_run = p.add_run(f"{turn.speaker.upper()}:   ")
                speaker_run.font.name = "Courier New"
            text_run = p.add_run(turn.text)
            text_run.font.name = "Courier New"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def parse_docx_to_turns(docx_bytes: bytes) -> List[dict]:
    """
    Parse a DOCX transcript into speaker/text turns.
    """
    buffer = io.BytesIO(docx_bytes)
    doc = Document(buffer)

    title_page_patterns = [
        r"^generated\s+transcript\s*$",
        r"^case\s+name:\s*",
        r"^case\s+number:\s*",
        r"^date:\s*",
        r"^time:\s*",
        r"^location:\s*",
        r"^original\s+file:\s*",
        r"^duration:\s*",
        r"^firm\s*(name|or\s+organization)?\s*:\s*",
    ]
    title_page_regex = re.compile("|".join(title_page_patterns), re.IGNORECASE)

    turns = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if title_page_regex.match(text):
            continue

        match = re.match(r"^([A-Z][A-Z0-9\s\-\.\']*?):\s{1,5}(.+)$", text, re.IGNORECASE)
        if match:
            speaker = match.group(1).strip().upper()
            content = match.group(2).strip()
            if speaker and content:
                turns.append(
                    {
                        "speaker": speaker,
                        "text": content,
                        "is_continuation": bool(turns and turns[-1]["speaker"] == speaker),
                    }
                )
        else:
            if turns and not text.startswith("["):
                turns.append(
                    {
                        "speaker": turns[-1]["speaker"],
                        "text": text,
                        "is_continuation": True,
                    }
                )
            elif text and not text.startswith("["):
                turns.append(
                    {
                        "speaker": "UNKNOWN",
                        "text": text,
                        "is_continuation": False,
                    }
                )

    logger.info("Parsed %d turns from DOCX (legacy parser)", len(turns))
    return turns
